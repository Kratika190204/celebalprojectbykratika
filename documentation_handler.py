import streamlit as st
import os
from docx import Document
import io

class DocumentationHandler:
    def __init__(self):
        self.doc_file = "Project_Report.docx"
    
    def create_sample_documentation(self):
        """Create a sample documentation file when the original is missing or empty"""
        try:
            # Create a new document
            doc = Document()
            
            # Add title
            title = doc.add_heading('CLV Model Interface - Project Report', 0)
            
            # Add overview section
            doc.add_heading('Overview', level=1)
            doc.add_paragraph('The Customer Lifetime Value (CLV) Predictor is a comprehensive tool designed to help businesses predict and analyze the long-term value of their customers using advanced machine learning techniques.')
            
            # Add features section
            doc.add_heading('Key Features', level=1)
            features = [
                'Data Input Options: Upload CSV files, generate sample data, or manually enter customer information',
                'Advanced Predictions: ML-powered CLV predictions with confidence intervals',
                'Customer Segmentation: Automatic segmentation into High Value, Medium Value, and Low Value customers',
                'Interactive Visualizations: Dynamic charts and graphs for data exploration',
                'Churn Risk Analysis: Predict which customers are at risk of churning',
                'Export Capabilities: Download results in CSV or Excel format'
            ]
            for feature in features:
                doc.add_paragraph(feature, style='List Bullet')
            
            # Add data requirements section
            doc.add_heading('Data Requirements', level=1)
            doc.add_paragraph('The model requires the following customer data fields:')
            requirements = [
                'customer_id: Unique identifier',
                'age: Customer age (18-100)',
                'total_purchases: Number of purchases made',
                'avg_order_value: Average order value in USD',
                'days_since_first_purchase: Days since first purchase',
                'days_since_last_purchase: Days since last purchase',
                'acquisition_channel: Marketing channel (Online, Store, Social Media, Referral)',
                'location: Geographic location (Urban, Suburban, Rural)',
                'subscription_status: Current status (Active, Inactive, None)'
            ]
            for req in requirements:
                doc.add_paragraph(req, style='List Bullet')
            
            # Add model parameters section
            doc.add_heading('Model Parameters', level=1)
            params = [
                'Time Horizon: Prediction period (6, 12, 24, or 36 months)',
                'Discount Rate: Annual discount rate for future cash flows (0-20%)',
                'Confidence Threshold: Minimum confidence level for predictions (50-95%)'
            ]
            for param in params:
                doc.add_paragraph(param, style='List Bullet')
            
            # Add customer segments section
            doc.add_heading('Customer Segments', level=1)
            segments = [
                'High Value: Top 20% of customers by CLV',
                'Medium Value: Middle 60% of customers by CLV',
                'Low Value: Bottom 20% of customers by CLV'
            ]
            for segment in segments:
                doc.add_paragraph(segment, style='List Bullet')
            
            # Add usage instructions
            doc.add_heading('Usage Instructions', level=1)
            instructions = [
                'Select your data source from the sidebar',
                'Configure prediction parameters',
                'Click "Generate CLV Predictions"',
                'Explore results using interactive visualizations',
                'Export predictions for further analysis'
            ]
            for instruction in instructions:
                doc.add_paragraph(instruction, style='List Number')
            
            # Add technical details
            doc.add_heading('Technical Details', level=1)
            doc.add_paragraph('The model uses ensemble machine learning techniques combining:')
            tech_details = [
                'Random Forest Regression',
                'Gradient Boosting',
                'Feature engineering for temporal and behavioral patterns',
                'Cross-validation for model reliability'
            ]
            for detail in tech_details:
                doc.add_paragraph(detail, style='List Bullet')
            
            # Add support section
            doc.add_heading('Support', level=1)
            doc.add_paragraph('For technical support or questions, contact: kratikasoni73@gmail.com')
            
            # Save the document
            doc.save(self.doc_file)
            st.success(f"✅ Created sample documentation file: {self.doc_file}")
            
        except Exception as e:
            st.error(f"Error creating documentation file: {str(e)}")
    
    def display_documentation(self):
        """Display documentation from Project_Report.docx"""
        st.subheader("📚 Project Documentation")
        
        # Check if file exists and has content
        if os.path.exists(self.doc_file):
            file_size = os.path.getsize(self.doc_file)
            if file_size == 0:
                st.warning("Documentation file is empty. Creating sample documentation...")
                self.create_sample_documentation()
                # Try to read again after creating
                if os.path.exists(self.doc_file) and os.path.getsize(self.doc_file) > 0:
                    pass  # Continue to read the file
                else:
                    st.error("Failed to create documentation file.")
                    self._display_fallback_documentation()
                    return
        
        if os.path.exists(self.doc_file) and os.path.getsize(self.doc_file) > 0:
            try:
                # Read the DOCX file
                doc = Document(self.doc_file)
                
                # Display document content
                st.markdown("### CLV Model Interface - Project Report")
                
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        # Check if it's a heading (bold text or starts with numbers)
                        if paragraph.style and paragraph.style.name and paragraph.style.name.startswith('Heading'):
                            level = int(paragraph.style.name.replace('Heading ', ''))
                            st.markdown(f"{'#' * (level + 2)} {paragraph.text}")
                        else:
                            st.markdown(paragraph.text)
                
                # Provide download option
                st.markdown("---")
                
                with open(self.doc_file, "rb") as file:
                    st.download_button(
                        label="📄 Download Full Documentation",
                        data=file.read(),
                        file_name=self.doc_file,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        type="primary"
                    )
                    
            except Exception as e:
                st.error(f"Error reading documentation file: {str(e)}")
                st.info("Creating a new documentation file...")
                self.create_sample_documentation()
                self._display_fallback_documentation()
        else:
            st.warning("Documentation file 'Project_Report.docx' not found or empty.")
            st.info("Creating sample documentation...")
            self.create_sample_documentation()
            self._display_fallback_documentation()
    
    def _display_fallback_documentation(self):
        """Display fallback documentation when file is not available"""
        st.markdown("""
        ### 📖 CLV Model Interface Documentation
        
        #### Overview
        The Customer Lifetime Value (CLV) Predictor is a comprehensive tool designed to help businesses 
        predict and analyze the long-term value of their customers using advanced machine learning techniques.
        
        #### Features
        - **Data Input Options**: Upload CSV files, generate sample data, or manually enter customer information
        - **Advanced Predictions**: ML-powered CLV predictions with confidence intervals
        - **Customer Segmentation**: Automatic segmentation into High Value, Medium Value, and Low Value customers
        - **Interactive Visualizations**: Dynamic charts and graphs for data exploration
        - **Churn Risk Analysis**: Predict which customers are at risk of churning
        - **Export Capabilities**: Download results in CSV or Excel format
        
        #### Data Requirements
        The model requires the following customer data fields:
        - **customer_id**: Unique identifier
        - **age**: Customer age (18-100)
        - **total_purchases**: Number of purchases made
        - **avg_order_value**: Average order value in USD
        - **days_since_first_purchase**: Days since first purchase
        - **days_since_last_purchase**: Days since last purchase
        - **acquisition_channel**: Marketing channel (Online, Store, Social Media, Referral)
        - **location**: Geographic location (Urban, Suburban, Rural)
        - **subscription_status**: Current status (Active, Inactive, None)
        
        #### Model Parameters
        - **Time Horizon**: Prediction period (6, 12, 24, or 36 months)
        - **Discount Rate**: Annual discount rate for future cash flows (0-20%)
        - **Confidence Threshold**: Minimum confidence level for predictions (50-95%)
        
        #### Customer Segments
        - **High Value**: Top 20% of customers by CLV
        - **Medium Value**: Middle 60% of customers by CLV
        - **Low Value**: Bottom 20% of customers by CLV
        
        #### Usage Instructions
        1. Select your data source from the sidebar
        2. Configure prediction parameters
        3. Click "Generate CLV Predictions"
        4. Explore results using interactive visualizations
        5. Export predictions for further analysis
        
        #### Technical Details
        The model uses ensemble machine learning techniques combining:
        - Random Forest Regression
        - Gradient Boosting
        - Feature engineering for temporal and behavioral patterns
        - Cross-validation for model reliability
        
        #### Support
        For technical support or questions, contact: kratikasoni73@gmail.com
        """)
        
        # Create a sample documentation file for download
        sample_doc = """
        CLV Model Interface - Project Report
        
        This is a comprehensive Customer Lifetime Value prediction tool that helps businesses
        understand and predict the long-term value of their customers.
        
        Key Features:
        - Advanced ML predictions
        - Customer segmentation
        - Interactive visualizations
        - Export capabilities
        - Real-time analysis
        
        For more information, contact: kratikasoni73@gmail.com
        """
        
        st.download_button(
            label="📄 Download Sample Documentation",
            data=sample_doc,
            file_name="CLV_Model_Documentation.txt",
            mime="text/plain",
            type="secondary"
        )